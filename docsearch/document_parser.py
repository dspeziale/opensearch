"""
DocSearch - Document Parser
Supporta: PDF, DOC, DOCX, Excel, HTML, Markdown, TXT, MSG, PST
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional
import logging

# PDF
import PyPDF2

# Word
from docx import Document

# Excel
import pandas as pd
import openpyxl

# HTML
from bs4 import BeautifulSoup

# Markdown
import markdown

# Outlook
try:
    import extract_msg
    MSG_AVAILABLE = True
except ImportError:
    MSG_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("extract-msg not available. .msg files will not be supported.")

try:
    import pypff
    PST_AVAILABLE = True
except ImportError:
    PST_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("pypff not available. .pst files will not be supported.")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser universale per documenti"""

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'PDF Document',
        '.doc': 'Word Document',
        '.docx': 'Word Document',
        '.xls': 'Excel Spreadsheet',
        '.xlsx': 'Excel Spreadsheet',
        '.html': 'HTML Document',
        '.htm': 'HTML Document',
        '.md': 'Markdown Document',
        '.txt': 'Text Document',
        '.csv': 'CSV File',
        '.msg': 'Outlook Message',
        '.pst': 'Outlook Archive'
    }

    def __init__(self):
        self.parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,  # python-docx supporta anche .doc
            '.xlsx': self._parse_excel,
            '.xls': self._parse_excel,
            '.csv': self._parse_csv,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.md': self._parse_markdown,
            '.txt': self._parse_text,
            '.msg': self._parse_msg,
            '.pst': self._parse_pst
        }

    def parse(self, file_path: str) -> Dict:
        """
        Parse un documento e ritorna metadati + contenuto

        Returns:
            {
                'filename': str,
                'extension': str,
                'type': str,
                'content': str,
                'metadata': dict,
                'success': bool,
                'error': str (optional)
            }
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return {
                'success': False,
                'error': f'File not found: {file_path}'
            }

        extension = file_path.suffix.lower()

        if extension not in self.parsers:
            return {
                'success': False,
                'error': f'Unsupported file type: {extension}'
            }

        try:
            # Parse del contenuto
            content = self.parsers[extension](file_path)

            # Metadata
            metadata = {
                'filename': file_path.name,
                'extension': extension,
                'type': self.SUPPORTED_EXTENSIONS.get(extension, 'Unknown'),
                'size': file_path.stat().st_size,
                'path': str(file_path.absolute())
            }

            # Estrai keywords e summary
            keywords = self._extract_keywords(content)
            summary = self._generate_summary(content)

            return {
                'success': True,
                'filename': file_path.name,
                'extension': extension,
                'type': metadata['type'],
                'content': content,
                'metadata': metadata,
                'keywords': keywords,
                'summary': summary
            }

        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'filename': file_path.name
            }

    def _parse_pdf(self, file_path: Path) -> str:
        """Estrae testo da PDF"""
        text = []

        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(f"\n--- Page {page_num + 1} ---\n")
                        text.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num}: {e}")

        return '\n'.join(text)

    def _parse_docx(self, file_path: Path) -> str:
        """Estrae testo da Word DOCX"""
        doc = Document(file_path)

        text = []

        # Paragrafi
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        # Tabelle
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells])
                text.append(row_text)

        return '\n'.join(text)

    def _parse_excel(self, file_path: Path) -> str:
        """Estrae testo da Excel"""
        text = []

        # Leggi tutti i fogli
        excel_file = pd.ExcelFile(file_path)

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            text.append(f"\n=== Sheet: {sheet_name} ===\n")

            # Headers
            headers = ' | '.join(str(col) for col in df.columns)
            text.append(headers)
            text.append('-' * len(headers))

            # Rows (max 100 per performance)
            for idx, row in df.head(100).iterrows():
                row_text = ' | '.join(str(val) for val in row.values)
                text.append(row_text)

        return '\n'.join(text)

    def _parse_csv(self, file_path: Path) -> str:
        """Estrae testo da CSV"""
        df = pd.read_csv(file_path)

        text = []

        # Headers
        headers = ' | '.join(str(col) for col in df.columns)
        text.append(headers)
        text.append('-' * len(headers))

        # Rows (max 100)
        for idx, row in df.head(100).iterrows():
            row_text = ' | '.join(str(val) for val in row.values)
            text.append(row_text)

        return '\n'.join(text)

    def _parse_html(self, file_path: Path) -> str:
        """Estrae testo da HTML"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'lxml')

        # Rimuovi script e style
        for script in soup(["script", "style"]):
            script.decompose()

        # Estrai testo
        text = soup.get_text()

        # Pulizia
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        return text

    def _parse_markdown(self, file_path: Path) -> str:
        """Estrae testo da Markdown"""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Converti a HTML poi a testo per avere output pulito
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()

        # Ma mantieni anche il markdown originale per contesto
        return f"=== MARKDOWN SOURCE ===\n{md_content}\n\n=== RENDERED TEXT ===\n{text}"

    def _parse_text(self, file_path: Path) -> str:
        """Estrae testo da file di testo"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _extract_keywords(self, text: str, max_keywords: int = 20) -> List[str]:
        """
        Estrae keywords dal testo usando semplice analisi di frequenza
        """
        # Rimuovi punteggiatura e converti a lowercase
        clean_text = re.sub(r'[^\w\s]', ' ', text.lower())

        # Split in parole
        words = clean_text.split()

        # Rimuovi stopwords comuni (semplice)
        stopwords = {
            'il', 'lo', 'la', 'i', 'gli', 'le', 'un', 'uno', 'una',
            'di', 'a', 'da', 'in', 'con', 'su', 'per', 'tra', 'fra',
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were', 'be', 'been',
            'e', 'è', 'che', 'del', 'della', 'dei', 'delle', 'nel', 'nella',
            'page', 'document', 'file', 'text'
        }

        # Filtra parole brevi e stopwords
        filtered_words = [
            word for word in words
            if len(word) > 3 and word not in stopwords
        ]

        # Conta frequenze
        word_freq = {}
        for word in filtered_words:
            word_freq[word] = word_freq.get(word, 0) + 1

        # Ordina per frequenza
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)

        # Prendi top N
        keywords = [word for word, freq in sorted_words[:max_keywords]]

        return keywords

    def _generate_summary(self, text: str, max_chars: int = 500) -> str:
        """
        Genera un breve summary del documento
        """
        # Prendi i primi N caratteri
        summary = text[:max_chars].strip()

        # Trova l'ultimo punto per non tagliare a metà frase
        last_period = summary.rfind('.')
        if last_period > 0:
            summary = summary[:last_period + 1]

        return summary

    def _parse_msg(self, file_path: Path) -> str:
        """Estrae testo da file Outlook MSG"""
        if not MSG_AVAILABLE:
            raise ImportError("extract-msg library not available. Install with: pip install extract-msg")

        text = []

        try:
            msg = extract_msg.Message(str(file_path))

            # Header
            text.append("=== EMAIL MESSAGE ===\n")

            # From
            if msg.sender:
                text.append(f"From: {msg.sender}")

            # To
            if msg.to:
                text.append(f"To: {msg.to}")

            # CC
            if msg.cc:
                text.append(f"CC: {msg.cc}")

            # Subject
            if msg.subject:
                text.append(f"Subject: {msg.subject}")

            # Date
            if msg.date:
                text.append(f"Date: {msg.date}")

            text.append("\n--- Body ---\n")

            # Body (prova prima HTML, poi plain text)
            if msg.htmlBody:
                # Parse HTML body
                soup = BeautifulSoup(msg.htmlBody, 'html.parser')
                body_text = soup.get_text()
                text.append(body_text)
            elif msg.body:
                text.append(msg.body)

            # Attachments info
            if msg.attachments:
                text.append(f"\n--- Attachments ({len(msg.attachments)}) ---")
                for i, attachment in enumerate(msg.attachments, 1):
                    att_name = getattr(attachment, 'longFilename', None) or getattr(attachment, 'shortFilename', 'unknown')
                    text.append(f"{i}. {att_name}")

            msg.close()

        except Exception as e:
            logger.error(f"Error parsing MSG file: {e}")
            raise

        return '\n'.join(text)

    def extract_msg_attachments(self, file_path: str, output_dir: str) -> List[Dict]:
        """
        Estrae gli allegati da un file .msg e li salva in una directory

        Args:
            file_path: Path al file .msg
            output_dir: Directory dove salvare gli allegati

        Returns:
            Lista di dict con info sugli allegati estratti:
            [
                {
                    'filename': str,
                    'path': str,
                    'size': int,
                    'success': bool,
                    'error': str (optional)
                }
            ]
        """
        if not MSG_AVAILABLE:
            logger.warning("extract-msg not available, cannot extract attachments")
            return []

        attachments_info = []
        file_path = Path(file_path)
        output_dir = Path(output_dir)

        # Crea directory output se non esiste
        output_dir.mkdir(parents=True, exist_ok=True)

        try:
            msg = extract_msg.Message(str(file_path))

            if not msg.attachments:
                logger.info(f"No attachments found in {file_path.name}")
                msg.close()
                return []

            logger.info(f"Found {len(msg.attachments)} attachments in {file_path.name}")

            for i, attachment in enumerate(msg.attachments):
                try:
                    # Get attachment filename
                    att_name = getattr(attachment, 'longFilename', None) or \
                               getattr(attachment, 'shortFilename', None) or \
                               f"attachment_{i+1}"

                    # Sanitize filename
                    att_name = "".join(c for c in att_name if c.isalnum() or c in '._- ')
                    att_name = att_name.strip()

                    # Get attachment data
                    att_data = attachment.data

                    if att_data:
                        # Save attachment
                        att_path = output_dir / att_name

                        # Se esiste già, aggiungi numero
                        counter = 1
                        original_path = att_path
                        while att_path.exists():
                            name_parts = original_path.stem, original_path.suffix
                            att_path = output_dir / f"{name_parts[0]}_{counter}{name_parts[1]}"
                            counter += 1

                        with open(att_path, 'wb') as f:
                            f.write(att_data)

                        attachments_info.append({
                            'filename': att_name,
                            'path': str(att_path),
                            'size': len(att_data),
                            'success': True
                        })

                        logger.info(f"✓ Extracted attachment: {att_name} ({len(att_data)} bytes)")
                    else:
                        logger.warning(f"✗ No data for attachment: {att_name}")
                        attachments_info.append({
                            'filename': att_name,
                            'path': '',
                            'size': 0,
                            'success': False,
                            'error': 'No data'
                        })

                except Exception as e:
                    logger.error(f"Error extracting attachment {i}: {e}")
                    attachments_info.append({
                        'filename': f"attachment_{i+1}",
                        'path': '',
                        'size': 0,
                        'success': False,
                        'error': str(e)
                    })

            msg.close()

        except Exception as e:
            logger.error(f"Error processing MSG file for attachments: {e}")
            return []

        return attachments_info

    def _parse_pst(self, file_path: Path) -> str:
        """Estrae testo da file Outlook PST (archivio)"""
        if not PST_AVAILABLE:
            raise ImportError("pypff library not available. Install with: pip install pypff")

        text = []

        try:
            pst = pypff.file()
            pst.open(str(file_path))

            text.append("=== OUTLOOK PST ARCHIVE ===\n")

            root = pst.get_root_folder()
            text.append(f"Root Folder: {root.name if hasattr(root, 'name') else 'PST Archive'}")
            text.append(f"Number of sub-folders: {root.get_number_of_sub_folders()}")

            # Funzione ricorsiva per processare folders
            def process_folder(folder, depth=0, max_emails=50):
                """Processa una cartella ricorsivamente (limitato a max_emails per performance)"""
                indent = "  " * depth
                folder_name = folder.name if hasattr(folder, 'name') else 'Unknown'

                text.append(f"\n{indent}📁 Folder: {folder_name}")

                # Processa messaggi nella cartella (max 50)
                emails_processed = 0
                for i in range(folder.get_number_of_sub_messages()):
                    if emails_processed >= max_emails:
                        text.append(f"{indent}  ... (more emails not shown)")
                        break

                    try:
                        message = folder.get_sub_message(i)

                        # Estrai info base
                        subject = message.get_subject() if hasattr(message, 'get_subject') else 'No Subject'
                        sender = message.get_sender_name() if hasattr(message, 'get_sender_name') else 'Unknown'

                        text.append(f"{indent}  📧 [{i+1}] {subject}")
                        text.append(f"{indent}      From: {sender}")

                        # Body (solo primi 500 caratteri per performance)
                        if hasattr(message, 'get_plain_text_body'):
                            body = message.get_plain_text_body()
                            if body:
                                body_preview = body[:500].strip()
                                text.append(f"{indent}      Body: {body_preview}...")

                        emails_processed += 1

                    except Exception as e:
                        logger.warning(f"Error processing message {i}: {e}")
                        continue

                # Processa sub-folders ricorsivamente
                for j in range(folder.get_number_of_sub_folders()):
                    try:
                        sub_folder = folder.get_sub_folder(j)
                        process_folder(sub_folder, depth + 1, max_emails)
                    except Exception as e:
                        logger.warning(f"Error processing sub-folder {j}: {e}")
                        continue

            # Processa tutto
            process_folder(root)

            pst.close()

        except Exception as e:
            logger.error(f"Error parsing PST file: {e}")
            raise

        return '\n'.join(text)


# Test
if __name__ == '__main__':
    parser = DocumentParser()

    # Test con un file markdown
    test_file = '../Documentazione/ai_soc_stack_guide.md'

    if os.path.exists(test_file):
        result = parser.parse(test_file)

        if result['success']:
            print(f"✅ Parsed: {result['filename']}")
            print(f"Type: {result['type']}")
            print(f"Content length: {len(result['content'])} chars")
            print(f"Keywords: {', '.join(result['keywords'][:10])}")
            print(f"\nSummary:\n{result['summary']}")
        else:
            print(f"❌ Error: {result['error']}")
    else:
        print(f"⚠️  Test file not found: {test_file}")
